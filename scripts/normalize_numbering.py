"""normalize_numbering.py — 统一编号连接符("-" 与 ".")。

关键修正:图/表 与 公式 是**两套独立的编号体例**,不可混为一谈。
- 图、表 通常用小数点:图3.1、表3.1。
- 公式(式/括号)通常用连字符:(3-1)、式(3-1)。
《写作指南》第16节"全文统一"指的是**同一编号序列内部一致**,而非把图表与
公式强行统一成同一连接符。故本脚本对两套体例分别统计、分别统一,默认
图/表→"."、公式→"-",且仅在某一套**自身混用**时才改动它。

这是整条流水线中唯一会改动文字的脚本。改动严格限定为编号连接符这一个字符,
每处替换都记入 numbering_changes.json,供逐条核对。

用法:
    统计(不改):
        python normalize_numbering.py 输入.docx --census
    分体例自动统一(图表→多数派/默认".",公式→多数派/默认"-"):
        python normalize_numbering.py 输入.docx --auto -o 输出.docx
    手动指定某一套(可单独或同时给出):
        python normalize_numbering.py 输入.docx --figtab dot --formula hyphen -o 输出.docx
    (兼容)旧版全文统一一种连接符——仅在确需违反分体例约定时使用:
        python normalize_numbering.py 输入.docx --unify-all hyphen|dot -o 输出.docx
"""
import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document

# 图/表 引导编号:图3.1 / 表3-1 / 续图2.1
PAT_FIGTAB = re.compile(r"(续?[图表])(\s*)([A-Z]|\d{1,2})([-.．])(\d{1,3})")
# 公式:式3-1 这种带"式"前缀的,以及 (3-1)/（3.1） 括号编号
PAT_FORMULA_LEAD = re.compile(r"(式)(\s*)([A-Z]|\d{1,2})([-.．])(\d{1,3})")
PAT_PAREN = re.compile(r"([（(])([A-Z]|\d{1,2})([-.．])(\d{1,3})([)）])")

CONNECTOR = {"hyphen": "-", "dot": "."}
KEY = {"-": "hyphen", ".": "dot", "．": "dot"}


def census_text(text, c_fig, c_form):
    for m in PAT_FIGTAB.finditer(text):
        c_fig[KEY[m.group(4)]] += 1
    for m in PAT_FORMULA_LEAD.finditer(text):
        c_form[KEY[m.group(4)]] += 1
    for m in PAT_PAREN.finditer(text):
        c_form[KEY[m.group(3)]] += 1


def retarget_text(t, tf_fig, tf_form):
    """把图表编号连接符改成 tf_fig、公式编号连接符改成 tf_form(都是单字符)。"""
    t = PAT_FIGTAB.sub(
        lambda m: m.group(0) if m.group(4) == tf_fig
        else f"{m.group(1)}{m.group(2)}{m.group(3)}{tf_fig}{m.group(5)}", t)
    t = PAT_FORMULA_LEAD.sub(
        lambda m: m.group(0) if m.group(4) == tf_form
        else f"{m.group(1)}{m.group(2)}{m.group(3)}{tf_form}{m.group(5)}", t)
    t = PAT_PAREN.sub(
        lambda m: m.group(0) if m.group(3) == tf_form
        else f"{m.group(1)}{m.group(2)}{tf_form}{m.group(4)}{m.group(5)}", t)
    return t


def remaining_mismatches(text, tf_fig, tf_form):
    bad = []
    for m in PAT_FIGTAB.finditer(text):
        if m.group(4) != tf_fig:
            bad.append(m.group(0))
    for m in PAT_FORMULA_LEAD.finditer(text):
        if m.group(4) != tf_form:
            bad.append(m.group(0))
    for m in PAT_PAREN.finditer(text):
        if m.group(3) != tf_form:
            bad.append(m.group(0))
    return bad


def replace_in_runs(p, tf_fig, tf_form, changes, loc):
    para_text = p.text
    if not remaining_mismatches(para_text, tf_fig, tf_form):
        return
    for run in p.runs:
        t = run.text
        if not t:
            continue
        new_t = retarget_text(t, tf_fig, tf_form)
        if new_t != t:
            n = sum(1 for a, b in zip(t, new_t) if a != b)
            changes.append({"loc": loc, "old": t.strip()[:50],
                            "new": new_t.strip()[:50], "替换字符数": n})
            run.text = new_t
    still = remaining_mismatches(p.text, tf_fig, tf_form)
    if still:
        changes.append({"loc": loc, "unresolved": still,
                        "note": "编号跨run拆分,未能自动替换,需复核"})


def iter_all_paragraphs(doc):
    for i, p in enumerate(doc.paragraphs):
        yield f"段{i}", p
    for ti, t in enumerate(doc.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    yield f"表{ti+1}({ri},{ci})", p


def pick(counter, default_key):
    """某一套体例的目标连接符:有多数派取多数派,平局/空集取默认。"""
    h, d = counter["hyphen"], counter["dot"]
    if h == 0 and d == 0:
        return default_key
    if h == d:
        return default_key
    return "hyphen" if h > d else "dot"


def report_line(name, counter):
    return (f"{name}: 连字符'-' {counter['hyphen']} 处, 小数点'.' {counter['dot']} 处"
            + ("  [自身混用!]" if counter["hyphen"] and counter["dot"] else "  [一致]"))


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    ap.add_argument("--census", action="store_true", help="只统计不修改")
    ap.add_argument("--auto", action="store_true",
                    help="分体例自动统一(图表默认'.',公式默认'-')")
    ap.add_argument("--figtab", choices=["hyphen", "dot"], help="指定图/表连接符")
    ap.add_argument("--formula", choices=["hyphen", "dot"], help="指定公式连接符")
    ap.add_argument("--unify-all", choices=["hyphen", "dot"],
                    help="兼容:图表与公式全部统一为同一连接符(一般不建议)")
    ap.add_argument("-o", "--out")
    ap.add_argument("-r", "--report", default="numbering_changes.json")
    args = ap.parse_args()

    doc = Document(args.docx)
    c_fig = {"hyphen": 0, "dot": 0}
    c_form = {"hyphen": 0, "dot": 0}
    for _, p in iter_all_paragraphs(doc):
        census_text(p.text, c_fig, c_form)
    print(report_line("图/表编号", c_fig))
    print(report_line("公式编号", c_form))

    if args.census:
        fig_mixed = c_fig["hyphen"] and c_fig["dot"]
        form_mixed = c_form["hyphen"] and c_form["dot"]
        if fig_mixed or form_mixed:
            print("某套体例自身混用,需要统一(图表与公式各自独立,互不影响)。")
        else:
            print("两套体例各自一致,无需处理。")
        return

    if not args.out:
        sys.exit("修改模式需要 -o")

    # 决定两套体例的目标连接符
    if args.unify_all:
        tf_fig_key = tf_form_key = args.unify_all
    else:
        # 默认:图表 dot,公式 hyphen;--auto 用多数派;显式 --figtab/--formula 覆盖
        tf_fig_key = args.figtab or (pick(c_fig, "dot") if args.auto else "dot")
        tf_form_key = args.formula or (pick(c_form, "hyphen") if args.auto else "hyphen")
    tf_fig, tf_form = CONNECTOR[tf_fig_key], CONNECTOR[tf_form_key]
    print(f"目标:图/表→'{tf_fig}'  公式→'{tf_form}'")

    changes = []
    for loc, p in iter_all_paragraphs(doc):
        replace_in_runs(p, tf_fig, tf_form, changes, loc)

    out = {"target_figtab": tf_fig, "target_formula": tf_form,
           "census_before": {"figtab": c_fig, "formula": c_form},
           "changes": changes}
    Path(args.report).write_text(json.dumps(out, ensure_ascii=False, indent=1),
                                 encoding="utf-8")
    doc.save(args.out)
    n_changed = sum(1 for c in changes if "old" in c)
    n_unres = sum(1 for c in changes if "unresolved" in c)
    if n_changed == 0 and n_unres == 0:
        print(f"两套体例已各自一致,未作任何改动 → {args.out}")
    else:
        print(f"已替换 {n_changed} 处,未解决 {n_unres} 处 → {args.out}")
    print(f"逐条记录 → {args.report}")
    if n_unres:
        print("[警告] 存在跨run编号,请按报告复核。")


if __name__ == "__main__":
    main()
