"""sections.py — 自动完成分节、篇眉、页码。零人工。

用法:
    python sections.py 输入.docx -m structure_map.json -o 输出.docx -r sections_report.json
        [--cn-cover IDX] [--en-cover IDX]

功能:
1. 在每个"部分"(章级标题段)起始处自动插入分节符:
   - 正文第一章:奇数页分节符(另页右页开始);其余:下一页分节符。
2. 每节篇眉 = 该部分章标题文字,五号(10.5pt),中文宋体/西文 TNR,居中,奇偶相同,首页同。
3. 页脚页码 PAGE 域,居中,TNR 五号,无修饰:
   - 摘要 → 符号和缩略语说明:大写罗马数字,从 Ⅰ 起连续;
   - 正文第一章 → 文末:阿拉伯数字,从 1 起连续。
4. 摘要之前的前置区(封面/英文封面/名单/授权):篇眉页脚清空,且各封面页与摘要
   均置于奇数页(另页右页)——这些都是单面印刷,双面打印时会自动在每页后补空白页
   (符合打印格式;原本已是奇数页则不重复补)。用 --no-front-blank 可关闭补空白页。
   --cn-cover/--en-cover 指定封面所在段的"部分序号"时,对该节应用封面专用页边距。
5. settings.xml 写入 updateFields,Word 打开时自动刷新目录/清单等所有域(免按 F9)。

必须在 apply.py 之后运行(二者都不增删段落,structure_map 始终有效)。
"""
import argparse
import copy
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

sys.path.insert(0, str(Path(__file__).parent))
import spec

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

FRONT_ROMAN_TITLES = {"摘要", "Abstract", "目录", "插图和附表清单", "插图清单",
                      "附表清单", "符号和缩略语说明"}
PART_START_ROLES = {"chapter_title", "abstract_en_title"}

CN_COVER_MARGINS = dict(top=6.0, bottom=6.0, left=4.0, right=4.0)
EN_COVER_MARGINS = dict(top=5.5, bottom=5.0, left=3.6, right=3.6)


def norm(t):
    return "".join(t.split()).replace("　", "")


def find_parts(smap_paragraphs):
    """返回 [(start_index, title_text, kind)], kind ∈ cover/front/body/back."""
    parts = []
    seen_abstract = False
    seen_body = False
    for e in smap_paragraphs:
        if e["role"] not in PART_START_ROLES:
            continue
        title = e["text"].strip()
        n = norm(title)
        if n.startswith("第") and "章" in n[:6]:
            kind = "body"
            seen_body = True
        elif n in FRONT_ROMAN_TITLES:
            kind = "front"
            seen_abstract = True
        elif n.startswith("附录"):
            kind = "body"
        elif not seen_abstract:
            kind = "cover"
        else:
            kind = "body" if not seen_body else "back"  # 引言等异名首章兜底为 body
        if n in ("参考文献", "致谢", "声明") or n.startswith("个人简历") \
                or "评语" in n or "决议书" in n:
            kind = "back"
        parts.append({"start": e["index"], "title": title, "kind": kind})
    return parts


_CN_DIGITS = "零一二三四五六七八九"


def _fmt_num(n, num_fmt):
    """把序号 n 按 Word numFmt 转成字符串(覆盖章/附录常见格式)。"""
    if num_fmt in ("upperLetter",):
        return _alpha(n).upper()
    if num_fmt in ("lowerLetter",):
        return _alpha(n).lower()
    if num_fmt in ("upperRoman", "lowerRoman"):
        r = _roman(n)
        return r.upper() if num_fmt == "upperRoman" else r.lower()
    if num_fmt in ("chineseCounting", "chineseCountingThousand",
                   "ideographTraditional", "ideographZodiac",
                   "japaneseCounting", "chineseLegalSimplified"):
        return _cn(n)
    # decimal / decimalZero / decimalFullWidth / 其它 → 阿拉伯数字
    return str(n)


def _alpha(n):
    s = ""
    while n > 0:
        n, r = divmod(n - 1, 26)
        s = chr(97 + r) + s
    return s or "a"


def _roman(n):
    table = [(1000, "m"), (900, "cm"), (500, "d"), (400, "cd"), (100, "c"),
             (90, "xc"), (50, "l"), (40, "xl"), (10, "x"), (9, "ix"),
             (5, "v"), (4, "iv"), (1, "i")]
    out = ""
    for v, s in table:
        while n >= v:
            out += s
            n -= v
    return out


def _cn(n):
    if n <= 0:
        return str(n)
    if n < 10:
        return _CN_DIGITS[n]
    if n == 10:
        return "十"
    if n < 20:
        return "十" + _CN_DIGITS[n - 10]
    if n < 100:
        t, o = divmod(n, 10)
        return _CN_DIGITS[t] + "十" + (_CN_DIGITS[o] if o else "")
    return str(n)


def _build_numbering_resolver(doc):
    """返回 resolve(numId) -> (numFmt, lvlText, start) | None,自动跟随
    numStyleLink / styleLink 链找到真正定义层级的 abstractNum。"""
    try:
        nb = doc.part.numbering_part.element
    except Exception:
        return lambda _x: None
    num2abs, abs_def, by_stylelink, abs_numstylelink = {}, {}, {}, {}
    for n in nb.findall(qn("w:num")):
        a = n.find(qn("w:abstractNumId"))
        if a is not None:
            num2abs[n.get(qn("w:numId"))] = a.get(qn("w:val"))
    for an in nb.findall(qn("w:abstractNum")):
        aid = an.get(qn("w:abstractNumId"))
        sl = an.find(qn("w:styleLink")); nsl = an.find(qn("w:numStyleLink"))
        if sl is not None:
            by_stylelink[sl.get(qn("w:val"))] = aid
        if nsl is not None:
            abs_numstylelink[aid] = nsl.get(qn("w:val"))
        for l in an.findall(qn("w:lvl")):
            if l.get(qn("w:ilvl")) == "0":
                ft = l.find(qn("w:numFmt")); tx = l.find(qn("w:lvlText"))
                stt = l.find(qn("w:start"))
                if ft is not None and tx is not None:
                    abs_def[aid] = (ft.get(qn("w:val")), tx.get(qn("w:val")),
                                    int(stt.get(qn("w:val"))) if stt is not None else 1)
    def resolve(num_id):
        a = num2abs.get(num_id)
        if a is None:
            return None
        if a in abs_def:
            return abs_def[a]
        s = abs_numstylelink.get(a)            # 跟随 numStyleLink → styleLink
        if s and s in by_stylelink and by_stylelink[s] in abs_def:
            return abs_def[by_stylelink[s]]
        return None
    return resolve


def _build_style_numid(doc):
    """返回 styleId -> 直接 numId(含 basedOn 继承)的解析函数。"""
    sp = doc.styles.element
    info = {}
    for s in sp.findall(qn("w:style")):
        sid = s.get(qn("w:styleId"))
        bo = s.find(qn("w:basedOn"))
        npr = s.find(qn("w:pPr") + "/" + qn("w:numPr") + "/" + qn("w:numId"))
        info[sid] = {
            "basedOn": bo.get(qn("w:val")) if bo is not None else None,
            "numId": npr.get(qn("w:val")) if npr is not None else None,
        }
    def style_numid(sid):
        seen = set()
        while sid and sid in info and sid not in seen:
            seen.add(sid)
            if info[sid]["numId"]:
                return info[sid]["numId"]
            sid = info[sid]["basedOn"]
        return None
    return style_numid


def _para_numid(p, style_numid):
    """段落的有效 numId:优先直接 numPr,否则取段落样式链上的 numId。"""
    npr = p._p.find(qn("w:pPr") + "/" + qn("w:numPr") + "/" + qn("w:numId"))
    if npr is not None:
        return npr.get(qn("w:val"))
    ps = p._p.find(qn("w:pPr") + "/" + qn("w:pStyle"))
    if ps is not None:
        return style_numid(ps.get(qn("w:val")))
    return None


def _has_literal_number(text, lvl_text):
    """正文标题文本是否已经手写了编号(如'第一章…'/'附录A…'),避免重复添加。
    依据 lvlText 模板(如 '第\xa0%1\xa0章'、'附录%1')构造正则,
    把空格/不间断空格当作可有可无。"""
    if "%1" not in lvl_text:
        return False
    pre, post = lvl_text.split("%1", 1)
    def lit(s):
        # 允许每个可见字符之间出现任意空白(含半角/全角/不间断空格)
        chars = [re.escape(c) for c in s.replace("\xa0", " ") if not c.isspace()]
        return r"\s*".join(chars)
    numtok = r"[0-9０-９A-Za-zⅠ-Ⅻⅰ-ⅻ一二三四五六七八九十百千零〇两]+"
    pat = r"^\s*" + lit(pre) + r"\s*" + numtok + r"\s*" + lit(post)
    return re.match(pat, text.replace("\xa0", " ")) is not None


def assign_chapter_headers(doc, parts):
    """为每个章级部分计算与正文标题完全一致的篇眉文本 part['header']:
      • 文本已手写编号(第一章…/附录A…) → 原样,不重复加;
      • 文本未写编号但段落自动编号(第N章 / 附录X 等) → 解析真实编号后前置,
        编号格式(decimal/upperLetter/chineseCounting…)与计数(各体例独立)
        均按文档实际定义还原;
      • 无编号(摘要/参考文献/致谢等) → 原文。
    始终保留原文的分隔空格,使篇眉与正文标题逐字一致。"""
    resolve = _build_numbering_resolver(doc)
    style_numid = _build_style_numid(doc)
    P = doc.paragraphs
    counters = {}                       # 每个 (numFmt,lvlText) 体例独立计数
    for part in parts:
        p = P[part["start"]]
        raw = p.text
        part.setdefault("header", raw)
        num_id = _para_numid(p, style_numid)
        info = resolve(num_id) if num_id else None
        if info is None:                # 无自动编号:原文(可能已含手写编号)
            part["header"] = raw
            continue
        num_fmt, lvl_text, start = info
        if _has_literal_number(raw, lvl_text):   # 已手写编号,勿重复
            part["header"] = raw
            continue
        key = (num_fmt, lvl_text)
        counters[key] = counters.get(key, start - 1) + 1
        prefix = lvl_text.replace("%1", _fmt_num(counters[key], num_fmt))
        part["header"] = prefix + raw


def clone_body_sectpr(doc):
    """克隆文档末尾 body 级 sectPr 作为插入模板,剥离页眉页脚引用与页码设置。"""
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    tpl = copy.deepcopy(sectPr)
    for tag in ("w:headerReference", "w:footerReference", "w:pgNumType", "w:type"):
        for el in tpl.findall(qn(tag)):
            tpl.remove(el)
    return tpl


def insert_section_breaks(doc, parts, report):
    """在每个部分起始段的前一段 pPr 中插入 sectPr。"""
    tpl = clone_body_sectpr(doc)
    paragraphs = doc.paragraphs
    for part in parts:
        i = part["start"]
        if i == 0:
            continue  # 文档开头无需分节
        prev = paragraphs[i - 1]
        pPr = prev._p.get_or_add_pPr()
        if pPr.find(qn("w:sectPr")) is not None:
            continue  # 已有分节符
        new_sect = copy.deepcopy(tpl)
        pPr.append(new_sect)
        report.append({"action": "插入分节符", "before_paragraph": i,
                       "part": part["title"]})


def set_pgnumtype(sectPr, fmt=None, start=None):
    el = sectPr.find(qn("w:pgNumType"))
    if el is None:
        el = sectPr.makeelement(qn("w:pgNumType"), {})
        sectPr.append(el)
    if fmt:
        el.set(qn("w:fmt"), fmt)
    if start is not None:
        el.set(qn("w:start"), str(start))
    elif el.get(qn("w:start")):
        del el.attrib[qn("w:start")]


def style_header_run(run):
    run.font.size = Pt(spec.HEADER_TEXT["size"])
    run.font.name = spec.F_TNR
    run._r.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), spec.F_SONG)


def clear_hdrftr(part):
    """彻底清空页眉/页脚:删除其下所有元素(含 w:sdt 内容控件、表格、段落),
    再加回一个空段落供写入。

    修复:Word "插入页码" 库会把页码域包进 <w:sdt> 内容控件;python-docx 的
    .paragraphs 不会进入 w:sdt,旧代码只删可见段落的 run,残留的 sdt 页码域会与
    新写入的 PAGE 域叠加,导致同一页出现两个页号。改为整体清空根元素即可避免。"""
    el = part._element  # <w:hdr> / <w:ftr>
    for child in list(el):
        el.remove(child)
    return part.add_paragraph()


def _set_header_bottom_border(p):
    """给页眉段落加底部横线(清华模板视觉风格,0.75 磅单线)。
    sz 单位 1/8 磅:sz=6 → 0.75pt(常用),sz=4 → 0.5pt。"""
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def write_header(section, text):
    hdr = section.header
    hdr.is_linked_to_previous = False
    p = clear_hdrftr(hdr)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        run = p.add_run(text)
        style_header_run(run)
        _set_header_bottom_border(p)   # 仅在有篇眉文字时画横线;封面等空篇眉不画


def write_footer_pagenum(section, empty=False):
    ftr = section.footer
    ftr.is_linked_to_previous = False
    p = clear_hdrftr(ftr)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if empty:
        return
    # PAGE 域: fldChar(begin) + instrText + fldChar(end)
    def fld_run(char_type=None, instr=None):
        r = p.add_run()
        r.font.size = Pt(spec.PAGE_NUMBER["size"])
        r.font.name = spec.F_TNR
        if char_type:
            fld = r._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): char_type})
            r._r.append(fld)
        if instr:
            it = r._r.makeelement(qn("w:instrText"), {})
            it.set(qn("xml:space"), "preserve")
            it.text = instr
            r._r.append(it)
        return r
    fld_run(char_type="begin")
    fld_run(instr=" PAGE ")
    fld_run(char_type="end")


def apply_cover_margins(section, margins, report, label):
    section.top_margin = Cm(margins["top"])
    section.bottom_margin = Cm(margins["bottom"])
    section.left_margin = Cm(margins["left"])
    section.right_margin = Cm(margins["right"])
    report.append({"action": f"应用{label}页边距",
                   "values": f"上{margins['top']}/下{margins['bottom']}"
                             f"/左{margins['left']}/右{margins['right']}cm"})


def enable_update_fields(doc, report):
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        el = settings.makeelement(qn("w:updateFields"), {qn("w:val"): "true"})
        settings.append(el)
    report.append({"action": "设置打开时自动更新域(updateFields)",
                   "note": "Word 首次打开会自动刷新目录/插图清单/附表清单页码,无需按F9"})


def disable_even_odd_and_titlepg(doc, sections):
    # 篇眉奇偶相同、首页同:确保未启用 evenAndOddHeaders / titlePg
    settings = doc.settings.element
    for el in settings.findall(qn("w:evenAndOddHeaders")):
        settings.remove(el)
    for sec in sections:
        for el in sec._sectPr.findall(qn("w:titlePg")):
            sec._sectPr.remove(el)


def section_nonempty(doc):
    """按文档顺序判断每个分节是否含**实际内容**(非空文字或图片)。
    返回 list[bool],长度 = 分节数。用于:只把"有内容"的封面分节置于奇数页,
    而文档里原有的空白分节(已经充当单面打印的空白页)保持原样,避免补出双倍空白。"""
    res = []
    has = False
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            if any((t.text or "").strip() for t in child.findall(".//" + qn("w:t"))) \
                    or child.find(".//" + qn("w:drawing")) is not None \
                    or child.find(".//" + qn("w:pict")) is not None:
                has = True
            if child.find(qn("w:pPr") + "/" + qn("w:sectPr")) is not None:
                res.append(has)
                has = False
        elif child.tag == qn("w:tbl"):
            if any((t.text or "").strip() for t in child.findall(".//" + qn("w:t"))) \
                    or child.find(".//" + qn("w:drawing")) is not None \
                    or child.find(".//" + qn("w:pict")) is not None:
                has = True
        elif child.tag == qn("w:sectPr"):     # body 末尾的最后一节
            res.append(has)
            has = False
    return res


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    ap.add_argument("-m", "--map", required=True)
    ap.add_argument("-o", "--out", required=True)
    ap.add_argument("-r", "--report", default="sections_report.json")
    ap.add_argument("--cn-cover", type=int, help="中文封面的部分序号(0起)")
    ap.add_argument("--en-cover", type=int, help="英文封面的部分序号(0起)")
    ap.add_argument("--no-front-blank", action="store_true",
                    help="不为前置封面区补单面打印空白页"
                         "(默认:各封面页与摘要置于奇数页/另页右页,双面打印时自动补空白页)")
    args = ap.parse_args()

    doc = Document(args.docx)
    smap = json.loads(Path(args.map).read_text(encoding="utf-8"))
    paragraphs_map = smap["paragraphs"] if isinstance(smap, dict) else smap
    report = []

    parts = find_parts(paragraphs_map)
    if not parts:
        sys.exit("结构图谱中未发现任何章级标题,无法分节")
    body_first = next((p for p in parts if p["kind"] == "body"), None)
    front_first = next((p for p in parts if p["kind"] == "front"), None)

    # 为自动编号的正文章补回"第N章"(篇眉需与正文标题一致)
    assign_chapter_headers(doc, parts)
    insert_section_breaks(doc, parts, report)

    # 分节后,doc.sections = 前置区(可有多节)+ 各部分,逐一对应。
    # 前置区(封面/名单/授权)在清华官方模板里常含多个分节(各有专用页边距),
    # 故 offset 不能假设为 1,应按实际反算:offset = 总节数 − 部分数。
    sections = doc.sections
    offset = len(sections) - len(parts)
    if offset < 0:
        sys.exit(f"分节数({len(sections)})少于部分数({len(parts)}),中止")

    # 清空全部前置区分节的篇眉页脚(封面等不显示页眉页码);并把各前置区分节
    # 置于奇数页(另页右页)——封面/英文封面/名单/授权说明都是单面印刷,双面
    # 打印时奇数页分节符会自动在每页后补一张空白页(若原本已是奇数页则不重复补)。
    front_blank = not args.no_front_blank
    nonempty = section_nonempty(doc) if front_blank else []
    for k in range(offset):
        # 仅把"有内容"的封面分节(封面/英文封面/名单/授权)置于奇数页;原有的空白
        # 分节保持原样,这样既能给缺空白页的文档补齐,又不会给已有空白页的文档补重。
        odd = k > 0 and front_blank and k < len(nonempty) and nonempty[k]
        if odd:
            sections[k].start_type = WD_SECTION.ODD_PAGE
        write_header(sections[k], "")
        write_footer_pagenum(sections[k], empty=True)
        report.append({"action": "前置区(摘要前)篇眉页脚清空"
                       + ("、置于奇数页(单面打印自动补空白页)" if odd else ""),
                       "section": k})

    roman_started = False
    arabic_started = False
    for pi, part in enumerate(parts):
        sec = sections[pi + offset]
        sectPr = sec._sectPr
        title_n = norm(part["title"])

        if part["kind"] == "cover":
            write_header(sec, "")
            write_footer_pagenum(sec, empty=True)
            if args.cn_cover == pi:
                apply_cover_margins(sec, CN_COVER_MARGINS, report, "中文封面")
            elif args.en_cover == pi:
                apply_cover_margins(sec, EN_COVER_MARGINS, report, "英文封面")
            report.append({"action": "封面/前置部分,无篇眉页码", "part": part["title"]})
            continue

        # 分节符类型:正文第一章另页右页(奇数页);摘要(双面印刷起点)也置于
        # 奇数页;其余各部分下一页即可。
        front_odd = part is front_first and front_blank
        sec.start_type = (WD_SECTION.ODD_PAGE if (part is body_first or front_odd)
                          else WD_SECTION.NEW_PAGE)
        # 页码
        if part["kind"] == "front":
            set_pgnumtype(sectPr, fmt="upperRoman",
                          start=1 if not roman_started else None)
            roman_started = True
            num_desc = "罗马数字" + ("(从Ⅰ起)" if part is front_first else "(连续)")
        else:
            set_pgnumtype(sectPr, fmt="decimal",
                          start=1 if not arabic_started else None)
            arabic_started = True
            num_desc = "阿拉伯数字" + ("(从1起)" if part is body_first else "(连续)")
        # 篇眉与页脚
        hdr_text = part.get("header", part["title"])
        write_header(sec, hdr_text)
        write_footer_pagenum(sec)
        report.append({"action": "设置篇眉与页码", "part": part["title"],
                       "篇眉": hdr_text, "页码": num_desc,
                       "分节类型": "奇数页" if part is body_first else "下一页"})

    disable_even_odd_and_titlepg(doc, sections)
    enable_update_fields(doc, report)

    doc.save(args.out)
    Path(args.report).write_text(json.dumps(report, ensure_ascii=False, indent=1),
                                 encoding="utf-8")
    print(f"分节 {len(sections)} 节,篇眉/页码已写入 → {args.out}")
    print(f"操作记录 → {args.report}")


if __name__ == "__main__":
    main()
