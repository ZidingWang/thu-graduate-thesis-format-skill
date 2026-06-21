"""classify.py — 解析论文 docx,逐段分类,输出结构图谱 structure_map.json.

用法:
    python classify.py 输入.docx [-o structure_map.json]

输出的 JSON 是一个列表,每项:
    {"index": 段落序号, "role": 角色, "text": 前60字预览, "via": 判定依据}

角色定义见 spec.ROLES。分类基于文本特征 + 原文档大纲级别/样式名的启发式,
**可能出错**——使用本 skill 的 Claude 必须人工复核图谱(尤其 equation /
fig_caption / table_source / 各部分边界),修正后再交给 apply.py。
"""
import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).parent))
import spec

RE_CHAPTER = re.compile(r"^第\s*\d+\s*章")
RE_SEC1 = re.compile(r"^\d+\.\d+(\s|　|[^\d.])")
RE_SEC2 = re.compile(r"^\d+\.\d+\.\d+(\s|　|[^\d.])")
RE_SEC3 = re.compile(r"^\d+\.\d+\.\d+\.\d+(\s|　|[^\d.])")
RE_APPENDIX = re.compile(r"^附录\s*[A-Z]")
RE_APPENDIX_SEC = re.compile(r"^[A-Z]\.\d+")
RE_FIG = re.compile(r"^(续?图|Fig\.?|Figure)\s*([A-Z]|\d+)[-.．]\d+")
RE_TABLE = re.compile(r"^(续?表|Table)\s*([A-Z]|\d+)[-.．]\d+")
RE_REF_ENTRY = re.compile(r"^\[\d+\]")
RE_EQ_NUM = re.compile(r"[(（]([A-Z]|\d+)[-.．]\d+[)）]\s*$")
RE_KEYWORDS = re.compile(r"^(关键词|Keywords?)\s*[:：]")
# 目录行:行尾"引导符 + 页码"。引导符可为 点串/省略号/制表符/多个空格(含全角),
# 页码为罗马数字或阿拉伯数字。很多论文目录用【制表符】对齐页码(如 "第1章 引言\t6"),
# 旧正则只认点串/省略号,会把这种目录行误判成真正的章标题,连锁导致分节/页码错乱。
RE_TOC_LINE = re.compile(
    r"(?:\.{4,}|…+|\t+|[ 　]{2,})[ 　]*[IVXLCDMivxlcdm\d]+[ 　]*$")
RE_SOURCE = re.compile(r"^(资料来源|数据来源|来源)\s*[:：]")

CHAPTER_LEVEL_SET = {t.replace(" ", "") for t in spec.CHAPTER_LEVEL_TITLES}


def has_math(paragraph) -> bool:
    return bool(paragraph._p.findall(
        ".//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath"))


def norm(text: str) -> str:
    return re.sub(r"[\s　]+", "", text)


def build_outline_ctx(doc):
    """返回 outline_of_style(styleId) -> 大纲级别(0=章/标题1, 1=节/标题2, …);
    兼容样式里显式的 w:outlineLvl、"Heading N"/"标题 N" 命名,以及 basedOn 继承。
    用于识别**靠样式/自动编号**而非"第N章"文字的章节标题。无法判定返回 None。"""
    sp = doc.styles.element
    info = {}
    for st in sp.findall(qn("w:style")):
        sid = st.get(qn("w:styleId"))
        bo = st.find(qn("w:basedOn"))
        olvl = st.find(qn("w:pPr") + "/" + qn("w:outlineLvl"))
        name_el = st.find(qn("w:name"))
        name = name_el.get(qn("w:val")) if name_el is not None else ""
        lvl = None
        if olvl is not None:
            try:
                lvl = int(olvl.get(qn("w:val")))
            except (TypeError, ValueError):
                lvl = None
        if lvl is None:
            m = re.search(r"(?:heading|标题)\s*([1-9])", name, re.I)
            if m:
                lvl = int(m.group(1)) - 1
        info[sid] = {"basedOn": bo.get(qn("w:val")) if bo is not None else None,
                     "lvl": lvl}

    def outline_of_style(sid):
        seen = set()
        while sid and sid in info and sid not in seen:
            seen.add(sid)
            if info[sid]["lvl"] is not None:
                return info[sid]["lvl"]
            sid = info[sid]["basedOn"]
        return None
    return outline_of_style


def para_outline(p, outline_of_style):
    """段落有效大纲级别:直接 w:outlineLvl 优先,否则取段落样式链;
    9(Word 里=正文/无大纲)与缺省都返回 None。"""
    lvl = None
    direct = p._p.find(qn("w:pPr") + "/" + qn("w:outlineLvl"))
    if direct is not None:
        try:
            lvl = int(direct.get(qn("w:val")))
        except (TypeError, ValueError):
            lvl = None
    if lvl is None and outline_of_style is not None:
        ps = p._p.find(qn("w:pPr") + "/" + qn("w:pStyle"))
        if ps is not None:
            lvl = outline_of_style(ps.get(qn("w:val")))
    if lvl is not None and lvl >= 9:
        return None
    return lvl


def classify_paragraph(p, state, outline_of_style=None):
    """返回 (role, via)。state 记录当前所在部分,影响判定。"""
    text = p.text.strip()
    n = norm(text)
    style = (p.style.name or "") if p.style else ""

    if not text and not has_math(p):
        return "skip", "空段落"

    # —— 目录区内:带引导线/行尾页码的行优先判为目录行,防止"第N章…1"误判 ——
    if state["part"] == "toc" and RE_TOC_LINE.search(text):
        if RE_CHAPTER.match(text) or norm(re.split(r"[.…]", text)[0]) in CHAPTER_LEVEL_SET:
            return "toc_chapter", "目录章行"
        return "toc_entry", "目录行"

    # —— 章级标题与部分切换 ——
    # 章标题应是短标题;正文里"……第2章重点介绍……"这类**长句**也以"第N章"开头,
    # 必须用长度/句末标点排除,否则会被误判成章标题、凭空多出几个"章"。
    if RE_CHAPTER.match(text) and len(text) < 40 \
            and not text.rstrip().endswith(("。", "；", ".", ";")):
        state["part"] = "body"
        return "chapter_title", "正则:第N章"
    if n in CHAPTER_LEVEL_SET or n.rstrip("(（一二三四五六七八九十)）") in CHAPTER_LEVEL_SET:
        state["part"] = {"摘要": "abstract_cn", "Abstract": "abstract_en",
                         "目录": "toc", "参考文献": "references",
                         "致谢": "back", "声明": "back",
                         }.get(n, "front" if state["part"] in
                               ("front", "abstract_cn", "abstract_en", "toc") else "back")
        if n == "Abstract":
            return "abstract_en_title", "标题:Abstract"
        if n in ("插图和附表清单", "插图清单", "附表清单"):
            state["part"] = "lists"
        elif n == "符号和缩略语说明":
            state["part"] = "symbols"  # 单列出来:其条目有专用版式,需完全保留
        elif n.startswith("个人简历"):
            state["part"] = "cv"       # 个人简历/在学学术成果:子标题/标签/成果条目混排,保留
        return "chapter_title", f"章级标题:{n}"
    if RE_APPENDIX.match(text) and len(text) < 60:
        state["part"] = "appendix"
        return "chapter_title", "附录标题"

    # —— 个人简历、在学期间完成的相关学术成果区:子标题(个人简历/在学期间…,居中)、
    # 类别标签(学术论文:/专利:/参与科研项目:)、参考文献式成果条目混排,模板已排好版,
    # 整段保留不动——否则居中子标题会被改成两端对齐+首行缩进、条目也被加首行缩进而错乱。
    if state["part"] == "cv":
        return "cv_entry", "个人简历/在学学术成果(保留版式,不动)"

    # —— 按大纲级别 / Heading 样式识别自动编号的章/节标题 ——
    # 标题文字里没有"第N章"/"N.M"、靠 Word 自动编号或"标题1/标题2"样式的情形:
    # 大纲级别 0 → 章标题(同时把前置/目录/清单/符号区切换为正文);1~3 → 各级节标题。
    # 必须放在"符号说明区"判定之前——否则符号区后面的自动编号章会被符号区吞掉,
    # 导致状态一直停在 symbols。行尾带页码(目录行)或过长段落不计入,避免误判。
    olvl = para_outline(p, outline_of_style)
    if olvl is not None and 0 <= olvl <= 3 and 0 < len(n) < 80 \
            and not RE_TOC_LINE.search(text):
        if olvl == 0 and state["part"] != "cover":
            if state["part"] in ("front", "abstract_cn", "abstract_en",
                                  "toc", "lists", "symbols"):
                state["part"] = "body"
            return "chapter_title", "大纲级别1/标题1样式(自动编号章)"
        if olvl >= 1 and state["part"] in ("body", "appendix"):
            return f"section{olvl}", f"大纲级别{olvl + 1}/标题{olvl + 1}样式(自动编号节)"

    # —— 符号和缩略语说明区:有专用版式(w:tabs 两列对齐/悬挂缩进),整段保留不动 ——
    # 优先于下面的 toc/section 判定,避免被当成 toc_entry 后被 apply 改乱缩进。
    if state["part"] == "symbols" or "符号和缩略语" in style:
        return "symbol_entry", "符号和缩略语说明(保留专用版式,不动)"

    # —— 节标题(正文/附录内)——
    if state["part"] in ("body", "appendix"):
        if RE_SEC3.match(text):
            return "section3", "正则:N.M.K.L"
        if RE_SEC2.match(text):
            return "section2", "正则:N.M.K"
        if RE_SEC1.match(text) and len(text) < 60:
            return "section1", "正则:N.M"
        if state["part"] == "appendix" and RE_APPENDIX_SEC.match(text) and len(text) < 60:
            return "section1", "附录节标题"

    # —— 图表题注、来源 ——
    # 注意:目录/插图清单/附表清单区里的"图 1.1 …页码""表 3.2 …页码"是**清单条目**,
    # 不是题注——必须排除,否则会被套上题注的段前12/段后6磅,把清单撑开、行距变大、
    # 挤到两页。这些区交给下面的 toc/lists 判定(toc_entry,固定行距、无段前段后)。
    if state["part"] not in ("toc", "lists"):
        if RE_FIG.match(text):
            return "fig_caption", "正则:图N-M"
        if RE_TABLE.match(text):
            return "table_caption", "正则:表N-M"
        if RE_SOURCE.match(text):
            return "table_source", "正则:资料来源"

    # —— 公式段:含 OMML 或以编号(N-M)结尾且文字极少 ——
    if has_math(p):
        return "equation", "含OMML公式"
    if RE_EQ_NUM.search(text) and len(re.sub(r"[^\u4e00-\u9fff]", "", text)) <= 4:
        return "equation", "行末公式编号"

    # —— 各部分内容 ——
    if state["part"] == "toc" or RE_TOC_LINE.search(text):
        if RE_CHAPTER.match(text) or norm(text.split(".")[0]) in CHAPTER_LEVEL_SET:
            return "toc_chapter", "目录章行"
        return "toc_entry", "目录行"
    if state["part"] == "lists":
        return "toc_entry", "清单/符号说明行"
    if state["part"] == "references":
        if RE_REF_ENTRY.match(text):
            return "reference_entry", "正则:[N]"
        return "reference_entry", "参考文献部分内段落"
    if RE_KEYWORDS.match(text):
        return "keywords", "关键词行"
    if state["part"] == "cover":
        return "cover", "封面/前置页(不动)"

    # 兜底:正文
    via = f"兜底正文(原样式:{style})" if style else "兜底正文"
    return "body", via


PAT_NUM = re.compile(r"(?:续?[图表]|式|[（(])\s*(?:[A-Z]|\d{1,2})([-.．])\d{1,3}")


def _is_image_cell(cell) -> bool:
    """单元格内是否嵌有图片/形状(w:drawing 或 w:pict)。"""
    el = cell._tc
    return el.find(f".//{qn('w:drawing')}") is not None \
        or el.find(f".//{qn('w:pict')}") is not None


def _classify_table(t) -> str:
    """区分两类表格:
      · **数据表**(three_line):单元格里是文字/数字,是"真正的表格" → 画三线表。
      · **图片排版表**(image_layout):只是为了把图/子图摆整齐而建的布局表,
        单元格里嵌的是图片,原本无框/透明 → apply 阶段清除全部边框。

    判据:数据表不会在单元格里嵌图,因此**只要表中含嵌入图片且占比可观**即判为
    图片排版表;**不再限制行列数**——子图网格常见 2×2 / 3×2 / "一行图一行图注"
    等多行结构,旧的 行≤4/列≤3 会把六联子图这类表漏判成数据表、误画黑框。
    名单/封面信息表等若无图,默认仍是 three_line,确需保留原样的请人工标 skip。"""
    cells = [c for row in t.rows for c in row.cells]
    if not cells:
        return "three_line"
    n_img = sum(1 for c in cells if _is_image_cell(c))
    if n_img == 0:
        return "three_line"                  # 纯文字/数字 → 数据表
    if n_img / len(cells) >= 0.2:            # 含图占比可观 → 图片排版表(透明无边框)
        return "image_layout"
    return "three_line"


def _cover_region_table_indices(doc, cover_end):
    """返回位于封面/前置区(摘要段之前)的表格序号集合。

    学校模板的封面是**用表格排版**的:论文题目、培养单位/专业/作者/导师信息表、
    成文日期、英文封面、答辩委员会名单等都嵌在表格里。这些表格必须**完全不动**
    ——一旦被当成数据表画三线、把单元格字号压成 11pt 居中,封面标题会缩成小字、
    信息表会出现黑框,整个模板就崩了。故摘要之前的所有表格一律标 skip。"""
    if cover_end is None:
        return set()
    cover, para_seen, tbl_idx, passed = set(), 0, 0, False
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            if para_seen >= cover_end:
                passed = True
            para_seen += 1
        elif child.tag == qn("w:tbl"):
            if not passed:
                cover.add(tbl_idx)
            tbl_idx += 1
    return cover


def classify_document(docx_path: str):
    doc = Document(docx_path)
    outline_of_style = build_outline_ctx(doc)   # 样式→大纲级别,识别自动编号章节
    state = {"part": "cover"}  # cover → front → body → references → appendix → back
    paragraphs = []
    census = {"hyphen": 0, "dot": 0}
    cover_end = None                      # 摘要段索引:之前都属封面/前置区
    for i, p in enumerate(doc.paragraphs):
        role, via = classify_paragraph(p, state, outline_of_style)
        if cover_end is None and state["part"] != "cover":
            cover_end = i
        paragraphs.append({
            "index": i,
            "role": role,
            "text": p.text.strip()[:60],
            "via": via,
        })
        for m in PAT_NUM.finditer(p.text):
            census["hyphen" if m.group(1) == "-" else "dot"] += 1
    # 表格:封面/前置区的模板表格一律 skip(完全不动);其余按数据表→three_line
    # (画三线)、含图占位表→image_layout(清边框)。仍需保留原样的请人工标 skip。
    cover_tbls = _cover_region_table_indices(doc, cover_end)
    tables = []
    for ti, t in enumerate(doc.tables):
        first_row = " | ".join(c.text.strip()[:12] for c in t.rows[0].cells[:4]) \
            if t.rows else ""
        role = "skip" if ti in cover_tbls else _classify_table(t)
        tables.append({"index": ti, "rows": len(t.rows),
                       "cols": len(t.columns) if t.rows else 0,
                       "preview": first_row,
                       "role": role})
        for row in t.rows:
            for cell in row.cells:
                for m in PAT_NUM.finditer(cell.text):
                    census["hyphen" if m.group(1) == "-" else "dot"] += 1
    return {"paragraphs": paragraphs, "tables": tables,
            "numbering_census": census}


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    ap.add_argument("-o", "--out", default="structure_map.json")
    args = ap.parse_args()
    result = classify_document(args.docx)
    Path(args.out).write_text(
        json.dumps(result, ensure_ascii=False, indent=1), encoding="utf-8")
    roles = {}
    for r in result["paragraphs"]:
        roles[r["role"]] = roles.get(r["role"], 0) + 1
    print(f"共 {len(result['paragraphs'])} 段, {len(result['tables'])} 个表格 → {args.out}")
    for k, v in sorted(roles.items(), key=lambda x: -x[1]):
        print(f"  {k:18s} {v}")
    c = result["numbering_census"]
    print(f"编号体例: '-' {c['hyphen']} 处, '.' {c['dot']} 处"
          + (" [警告] 混用,需运行 normalize_numbering.py 统一"
             if c["hyphen"] and c["dot"] else ""))


if __name__ == "__main__":
    main()
