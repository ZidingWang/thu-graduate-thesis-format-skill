"""detect_degree.py — 自动识别研究生学位论文的学位类型,输出 degree_info.json.

用法:
    python detect_degree.py 论文.docx [-o degree_info.json]

判定依据(权威来源:《清华大学研究生学位论文写作指南》2025年3月版的封面/英文封面体例):
  - 中文封面学位类别行:"(申请清华大学……博士学位论文)" / "(……硕士学位论文)";
    专业学位为"……博士专业学位论文" / "……硕士专业学位论文"。
  - 英文封面学位行:"for the degree of Doctor of Philosophy" / "Master of …";
    以及 "Dissertation submitted to …"(博士)对 "Thesis submitted to …"(硕士)。
  - 关于学位论文使用授权的说明:博士/硕士两版,标题相同("关于学位论文使用授权的说明"),
    标题本身不可区分,正文细微差异,故仅作辅助、不作主判据。

判定范围:优先在"前置部分"(中文摘要之前的封面/名单/授权页)内检索这些标记,
避免正文内容里偶然出现的"博士/master"等词造成误判;前置部分无标记时再回退全文,
并降低置信度。博士与硕士的正文格式规范完全相同,本脚本只识别学位标识,不改动任何内容。
"""
import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document

# 与章标题同级的"前置部分结束"标志:命中第一个摘要/Abstract 视为正文前置区结束
RE_FRONT_END = re.compile(r"^(摘\s*要|Abstract)$", re.IGNORECASE)

# 标记表:(正则, 学位档次, 学位类别, 权重, 说明)
#   level: "doctoral" / "master" / None(仅判类别)
#   category: "professional" / "academic" / None
DEGREE_MARKERS = [
    # —— 中文封面:学位类别行(最强判据) ——
    (re.compile(r"博士专业学位论文"), "doctoral", "professional", 6, "中文封面:博士专业学位论文"),
    (re.compile(r"硕士专业学位论文"), "master", "professional", 6, "中文封面:硕士专业学位论文"),
    (re.compile(r"博士学位论文"), "doctoral", None, 5, "中文封面:博士学位论文"),
    (re.compile(r"硕士学位论文"), "master", None, 5, "中文封面:硕士学位论文"),
    # —— 英文封面:学位行 ——
    (re.compile(r"Doctor\s+of\s+Philosophy", re.I), "doctoral", "academic", 5, "英文封面:Doctor of Philosophy"),
    (re.compile(r"Doctoral\s+Dissertation", re.I), "doctoral", None, 4, "英文封面:Doctoral Dissertation"),
    (re.compile(r"Master\s+of\s+[A-Z][a-z]+", ), "master", None, 5, "英文封面:Master of …"),
    (re.compile(r"Master['’]?s?\s+(?:Degree|Thesis)", re.I), "master", None, 4, "英文封面:Master's Thesis/Degree"),
    (re.compile(r"Dissertation\s+submitted\s+to", re.I), "doctoral", None, 2, "英文封面:Dissertation submitted to(博士)"),
    (re.compile(r"Thesis\s+submitted\s+to", re.I), "master", None, 2, "英文封面:Thesis submitted to(硕士)"),
    # —— 学位类别(专业学位 / 学术学位)辅助标记 ——
    (re.compile(r"专业学位"), None, "professional", 0, "出现“专业学位”字样"),
    (re.compile(r"工程博士"), "doctoral", "professional", 3, "工程博士(专业学位)"),
    (re.compile(r"工程硕士"), "master", "professional", 3, "工程硕士(专业学位)"),
    # —— 仅在前置区出现的孤立“博士/硕士”词(弱判据) ——
    (re.compile(r"博\s*士"), "doctoral", None, 1, "前置区出现“博士”"),
    (re.compile(r"硕\s*士"), "master", None, 1, "前置区出现“硕士”"),
]


def all_lines(doc):
    """返回 (来源, 段索引, 文本) 列表;含段落与表格单元格(封面常用表格排版)。"""
    out = []
    for i, p in enumerate(doc.paragraphs):
        out.append(("para", i, p.text))
    for ti, t in enumerate(doc.tables):
        for row in t.rows:
            for cell in row.cells:
                out.append(("table", ti, cell.text))
    return out


def front_matter_cutoff(doc):
    """返回首个“摘要/Abstract”标题所在段索引;找不到返回 None。"""
    for i, p in enumerate(doc.paragraphs):
        if RE_FRONT_END.match(re.sub(r"[\s　]+", "", p.text)):
            return i
    return None


# 作者姓名字段标签(封面):申请人 / 研究生 / 作者 / 姓名
RE_NAME_LABEL = re.compile(r"^\s*(申\s*请\s*人|研\s*究\s*生|作\s*者|姓\s*名)\s*[：:]\s*(.*)$")
RE_CJK_NAME = re.compile(r"^[一-鿿·]{2,6}$")


def _clean_name(s):
    return re.sub(r"[\s　:：]+", "", s or "")


def extract_author(front_lines):
    """从前置区(封面)抓取作者姓名。支持两种模板存法:
      · 同段/同格:"研究生　：李 欣 雨" —— 冒号后即姓名;
      · 分格:封面信息表里 "申请人：" 单独一格,姓名 "潘 俊 生" 在相邻格。
    姓名常带全/半角空格(潘 俊 生),去空格后取 2~6 个汉字。排除指导教师/导师/
    作者签名等字段,以及授权说明里出现"研究生"的长句。找不到返回 None。"""
    for i, (src, idx, text) in enumerate(front_lines):
        t = (text or "").strip()
        if not t or len(t) > 25:                      # 跳过长句(授权说明等)
            continue
        if any(k in t for k in ("签名", "导师", "指导")):
            continue
        m = RE_NAME_LABEL.match(t)
        if not m:
            continue
        inline = _clean_name(m.group(2))
        if RE_CJK_NAME.match(inline):                 # 同段就有姓名
            return inline
        for j in range(i + 1, min(i + 4, len(front_lines))):  # 否则看相邻几格
            cand = _clean_name(front_lines[j][2])
            if RE_CJK_NAME.match(cand) and not any(
                    k in cand for k in ("签名", "导师", "日期", "学院", "专业", "大学")):
                return cand
    return None


def scan(lines, markers):
    """对给定文本行扫描标记,返回证据列表与分数。"""
    doc_score = {"doctoral": 0, "master": 0}
    cat_score = {"academic": 0, "professional": 0}
    evidence = []
    for src, idx, text in lines:
        if not text.strip():
            continue
        for pat, level, category, weight, label in markers:
            m = pat.search(text)
            if not m:
                continue
            if level and weight:
                doc_score[level] += weight
            if category and weight:
                cat_score[category] += weight
            evidence.append({
                "signal": label,
                "matched": m.group(0),
                "level": level,
                "category": category,
                "weight": weight,
                "source": src,
                "index": idx,
                "context": text.strip()[:80],
            })
    return doc_score, cat_score, evidence


def detect(docx_path):
    doc = Document(docx_path)
    lines = all_lines(doc)
    cutoff = front_matter_cutoff(doc)

    # 前置区:摘要标题之前的段落 + 全部表格(封面/名单/授权多用表格排版)
    if cutoff is not None:
        front = [ln for ln in lines if ln[0] == "table" or (ln[0] == "para" and ln[1] < cutoff)]
        front_detected = cutoff > 0
    else:
        # 无摘要标题:取前 40 段 + 表格作为前置区
        front = [ln for ln in lines if ln[0] == "table" or (ln[0] == "para" and ln[1] < 40)]
        front_detected = False

    author = extract_author(front)        # 从封面抓作者姓名(用于命名输出文件)

    doc_score, cat_score, evidence = scan(front, DEGREE_MARKERS)
    scope = "front_matter"

    # 前置区无任何学位档次判据 → 回退全文,降低置信度
    if doc_score["doctoral"] == 0 and doc_score["master"] == 0:
        doc_score, cat_score, evidence = scan(lines, DEGREE_MARKERS)
        scope = "whole_document"

    # —— 判定学位档次 ——
    d, m = doc_score["doctoral"], doc_score["master"]
    if d == 0 and m == 0:
        level, level_zh = "unknown", "未判定"
    elif d >= m:
        level, level_zh = "doctoral", "博士"
    else:
        level, level_zh = "master", "硕士"

    # —— 判定学位类别(专业 / 学术) ——
    if cat_score["professional"] > 0:
        category, category_zh = "professional", "专业学位"
    elif level != "unknown":
        category, category_zh = "academic", "学术学位"  # 默认学术学位
    else:
        category, category_zh = "unknown", "未判定"

    # —— 置信度 ——
    strong = any(e["weight"] >= 5 for e in evidence if e["level"] == level)
    conflict = d > 0 and m > 0 and min(d, m) / max(d, m) > 0.5
    if level == "unknown":
        confidence = "none"
    elif conflict:
        confidence = "low"
    elif strong and scope == "front_matter":
        confidence = "high"
    elif strong:
        confidence = "medium"
    else:
        confidence = "low"

    notes = []
    if not front_detected:
        notes.append("未在文档中找到独立的封面/前置区(可能封面单独成文件),判定基于全文检索,请人工确认。")
    if scope == "whole_document":
        notes.append("前置区未发现学位标记,已回退到全文检索,置信度下降。")
    if conflict:
        notes.append(f"博士与硕士标记同时较多(博士{d}/硕士{m}),可能是混排或引用,请人工确认。")
    if category == "academic" and cat_score["professional"] == 0 and level != "unknown":
        notes.append("未发现“专业学位”标记,默认按学术学位处理;若为专业学位请人工修正。")
    notes.append("博硕正文格式规范完全相同,学位类型仅影响封面与学位标识,不影响本 skill 的正文整改。")

    # —— 建议的输出文件名:姓名 + 博士/硕士论文 + 修改后版本 ——
    deg_word = {"doctoral": "博士论文", "master": "硕士论文"}.get(level, "研究生学位论文")
    if author:
        suggested_filename = f"{author}{deg_word}修改后版本"
    else:
        suggested_filename = None
        notes.append("未从封面抓到作者姓名(可手工命名);姓名一般在封面“申请人/研究生”栏或个人简历页。")

    return {
        "degree_level": level,
        "degree_level_zh": level_zh,
        "degree_category": category,
        "degree_category_zh": category_zh,
        "author_name": author,
        "suggested_filename": suggested_filename,
        "confidence": confidence,
        "scope": scope,
        "front_matter_detected": front_detected,
        "scores": {"doctoral": d, "master": m,
                   "academic": cat_score["academic"], "professional": cat_score["professional"]},
        "evidence": evidence,
        "notes": notes,
    }


def main():
    ap = argparse.ArgumentParser(description="识别研究生学位论文的学位类型(博士/硕士),输出 degree_info.json")
    ap.add_argument("docx")
    ap.add_argument("-o", "--out", default="degree_info.json")
    args = ap.parse_args()

    info = detect(args.docx)
    Path(args.out).write_text(
        json.dumps(info, ensure_ascii=False, indent=1), encoding="utf-8")

    print(f"学位档次: {info['degree_level_zh']} ({info['degree_level']})")
    print(f"学位类别: {info['degree_category_zh']} ({info['degree_category']})")
    print(f"作者姓名: {info.get('author_name') or '(未抓到,需人工命名)'}")
    if info.get("suggested_filename"):
        print(f"建议文件名: {info['suggested_filename']}.docx")
    print(f"置信度  : {info['confidence']}  (博士{info['scores']['doctoral']} / 硕士{info['scores']['master']}, 检索范围 {info['scope']})")
    if info["evidence"]:
        print("关键证据:")
        seen = set()
        for e in info["evidence"]:
            if e["signal"] in seen:
                continue
            seen.add(e["signal"])
            print(f"  - {e['signal']}  →  「{e['context']}」")
    for n in info["notes"]:
        print(f"注: {n}")
    print(f"→ {args.out}")


if __name__ == "__main__":
    main()
