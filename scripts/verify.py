"""verify.py — 终验:对比原始 docx 与最终 docx 的全文文本。

用法:
    python verify.py 原始.docx 最终.docx [--allow-numbering]

- 默认要求逐字相同(PASS/FAIL)。
- --allow-numbering:允许且仅允许图/表/式编号连接符("-"↔".")差异
  (配合 normalize_numbering.py 使用),其余任何差异都 FAIL 并打印 diff。
"""
import argparse
import difflib
import re
import sys

from docx import Document

PAT_NUM_CONN = re.compile(
    r"((?:续?[图表]|式|[（(])\s*(?:[A-Z]|\d{1,2}))[-.．](\d{1,3})")


def all_text_lines(path):
    doc = Document(path)
    lines = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                lines.append(cell.text)
    return lines


def canon(line):
    """把编号连接符归一为 '-',用于 --allow-numbering 比较。"""
    return PAT_NUM_CONN.sub(r"\1-\2", line)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("original")
    ap.add_argument("final")
    ap.add_argument("--allow-numbering", action="store_true")
    args = ap.parse_args()

    a = all_text_lines(args.original)
    b = all_text_lines(args.final)
    if args.allow_numbering:
        a_c = [canon(x) for x in a]
        b_c = [canon(x) for x in b]
    else:
        a_c, b_c = a, b

    if a_c == b_c:
        n_num = sum(1 for x, y in zip(a, b) if x != y)
        print("内容完整性终验: PASS")
        if args.allow_numbering and n_num:
            print(f"(其中 {n_num} 行仅编号连接符不同,属授权修改)")
        return

    print("内容完整性终验: FAIL —— 发现编号之外的文字差异:")
    diff = difflib.unified_diff(a_c, b_c, "原始", "最终", lineterm="", n=0)
    shown = 0
    for line in diff:
        if line.startswith(("+", "-")) and not line.startswith(("+++", "---")):
            print(" ", line[:120])
            shown += 1
            if shown >= 40:
                print("  …(差异过多,仅显示前40条)")
                break
    sys.exit(1)


if __name__ == "__main__":
    main()
