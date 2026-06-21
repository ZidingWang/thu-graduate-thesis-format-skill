"""apply.py — 按 2025 版规范整改格式。只改格式,绝不改文字。

用法:
    检查模式(不写文件,仅输出问题清单):
        python apply.py 论文.docx -m structure_map.json --dry-run -r report.json
    整改模式:
        python apply.py 论文.docx -m structure_map.json -o 论文_整改.docx -r report.json

report.json 记录每一处差异/修改:
    {"paragraph": 序号, "text": 预览, "role": 角色,
     "changes": [{"prop": 属性, "old": 旧值, "new": 新值}]}

设计约束:
- 永不修改 run.text / 段落文本 / 表格单元格文本。
- 字体修改按 中文(eastAsia)/西文(ascii,hAnsi) 分别设置。
- 不碰 run 的加粗/斜体/上下标(标题除外:章节标题统一不强制加粗,黑体本身即粗体效果)。
- 含 OMML 公式的 run 不改字体(只调段落属性),避免破坏公式渲染。
"""
import argparse
import copy
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

sys.path.insert(0, str(Path(__file__).parent))
import spec

ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}
ALIGN_NAME = {v: k for k, v in ALIGN.items()}

NO_TOUCH_ROLES = {"cover", "skip"}


def run_has_math(run) -> bool:
    return bool(run._r.findall(
        ".//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath"))


def para_has_math(p) -> bool:
    return bool(p._p.findall(
        ".//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath"))


def get_run_fonts(run):
    """返回 (ascii, eastAsia, size_pt)。读取 run 直接格式,不解析样式继承。"""
    rPr = run._r.rPr
    ascii_f = ea_f = None
    if rPr is not None and rPr.rFonts is not None:
        ascii_f = rPr.rFonts.get(qn("w:ascii"))
        ea_f = rPr.rFonts.get(qn("w:eastAsia"))
    size = run.font.size.pt if run.font.size else None
    return ascii_f, ea_f, size


# ────────────────────────────────────────────────────────────────────
# 字体继承解析:只有"解析后的有效字体 ≠ 规范"时才改 run,避免把本来
# 靠样式/文档默认值继承就正确的字体硬写到每一个 run 上(那样会让修改
# 数虚高、并把直接格式写满全文)。解析顺序遵循 OOXML:
#   run 直接 rPr → run 字符样式(rStyle)链 → 段落样式(pStyle)链
#   → docDefaults;主题字体(major/minor*)再经 theme1.xml 还原。
# 解析不确定(主题缺失/链中断)时返回 UNKNOWN → 按"不相等"处理 → 照常
# 写入,确保绝不因解析失误而漏改(合规性优先,数字下降其次)。
# ────────────────────────────────────────────────────────────────────
_FONT_CTX = None
_FONT_UNKNOWN = "<?>"


def _rfonts_from_rpr(rpr_el):
    d = {"ascii": None, "hAnsi": None, "eastAsia": None,
         "asciiTheme": None, "hAnsiTheme": None, "eastAsiaTheme": None, "sz": None}
    if rpr_el is None:
        return d
    rf = rpr_el.find(qn("w:rFonts"))
    if rf is not None:
        for k in ("ascii", "hAnsi", "eastAsia",
                  "asciiTheme", "hAnsiTheme", "eastAsiaTheme"):
            d[k] = rf.get(qn("w:" + k))
    sz = rpr_el.find(qn("w:sz"))
    if sz is not None:
        try:
            d["sz"] = int(sz.get(qn("w:val"))) / 2
        except (TypeError, ValueError):
            pass
    return d


def _ind_from_ind_el(ind_el):
    """从一个 <w:ind> 元素读取缩进属性(twips 或 1/100 字符,数值仅用于判断是否为 0)。"""
    d = {}
    if ind_el is None:
        return d
    for k in ("firstLine", "firstLineChars", "left", "leftChars",
              "hanging", "hangingChars", "start", "startChars"):
        v = ind_el.get(qn("w:" + k))
        if v is not None:
            try:
                d[k] = int(v)
            except ValueError:
                pass
    return d


def build_font_context(doc):
    ctx = {"pstyle": {}, "cstyle": {}, "docdef": None,
           "default_pstyle": None, "theme": {}}
    sp = doc.styles.element
    for st in sp.findall(qn("w:style")):
        sid = st.get(qn("w:styleId"))
        typ = st.get(qn("w:type"))
        bo = st.find(qn("w:basedOn"))
        d = _rfonts_from_rpr(st.find(qn("w:rPr")))
        d["basedOn"] = bo.get(qn("w:val")) if bo is not None else None
        d["ind"] = _ind_from_ind_el(st.find(qn("w:pPr") + "/" + qn("w:ind")))
        if typ == "character":
            ctx["cstyle"][sid] = d
        else:
            ctx["pstyle"][sid] = d
            if typ == "paragraph" and st.get(qn("w:default")) == "1":
                ctx["default_pstyle"] = sid
    dd = sp.find(qn("w:docDefaults") + "/" + qn("w:rPrDefault") + "/" + qn("w:rPr"))
    ctx["docdef"] = _rfonts_from_rpr(dd)
    ctx["docdef"]["basedOn"] = None
    ctx["docdef_ind"] = _ind_from_ind_el(
        sp.find(qn("w:docDefaults") + "/" + qn("w:pPrDefault")
                + "/" + qn("w:pPr") + "/" + qn("w:ind")))
    try:
        from lxml import etree
        A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
        for rel in doc.part.rels.values():
            if rel.reltype.endswith("/theme"):
                th = etree.fromstring(rel.target_part.blob)
                fs = th.find(".//" + A + "fontScheme")
                if fs is not None:
                    def face(mm, script):
                        f = fs.find(A + mm + "/" + A + script)
                        v = f.get("typeface") if f is not None else None
                        return v or None
                    ctx["theme"] = {
                        "majorLatin": face("majorFont", "latin"),
                        "minorLatin": face("minorFont", "latin"),
                        "majorEA": face("majorFont", "ea"),
                        "minorEA": face("minorFont", "ea"),
                    }
                break
    except Exception:
        pass
    return ctx


def _resolve_theme_token(token, ctx):
    th = ctx.get("theme", {})
    table = {
        "minorHAnsi": th.get("minorLatin"), "majorHAnsi": th.get("majorLatin"),
        "minorAscii": th.get("minorLatin"), "majorAscii": th.get("majorLatin"),
        "minorBidi": th.get("minorLatin"), "majorBidi": th.get("majorLatin"),
        "minorEastAsia": th.get("minorEA"), "majorEastAsia": th.get("majorEA"),
    }
    return table.get(token) or None


def _style_chain(start_id, table):
    out, seen, sid = [], set(), start_id
    while sid and sid in table and sid not in seen:
        seen.add(sid)
        out.append(table[sid])
        sid = table[sid].get("basedOn")
    return out


def effective_indent(p, ctx):
    """解析段落的有效缩进(随 段落直接 → 段落样式链 → docDefaults 继承)。
    返回 (first_line, hanging, left),任一非 0 即说明该段实际有缩进——
    用于判断"居中段落是否被(继承来的)首行/左缩进顶得右偏"。"""
    sources = [_ind_from_ind_el(p._p.find(qn("w:pPr") + "/" + qn("w:ind")))]
    ps = p._p.find(qn("w:pPr") + "/" + qn("w:pStyle"))
    pid = ps.get(qn("w:val")) if ps is not None else ctx.get("default_pstyle")
    for d in _style_chain(pid, ctx.get("pstyle", {})):
        sources.append(d.get("ind", {}))
    sources.append(ctx.get("docdef_ind", {}))

    def pick(keys):
        for src in sources:                 # 最近的来源优先(直接 > 样式 > 默认)
            for k in keys:
                if k in src:
                    return src[k]
        return 0
    first_line = pick(("firstLineChars", "firstLine"))
    hanging = pick(("hangingChars", "hanging"))
    left = pick(("leftChars", "left", "startChars", "start"))
    return first_line, hanging, left


def effective_run_fonts(run, p, ctx):
    """返回 (ascii_eff, eastAsia_eff, size_pt_eff);无法确定的项为 _FONT_UNKNOWN。"""
    sources = []
    rPr = run._r.rPr
    sources.append(_rfonts_from_rpr(rPr))
    if rPr is not None:
        rs = rPr.find(qn("w:rStyle"))
        if rs is not None:
            sources += _style_chain(rs.get(qn("w:val")), ctx["cstyle"])
    ps = p._p.find(qn("w:pPr") + "/" + qn("w:pStyle"))
    pid = ps.get(qn("w:val")) if ps is not None else ctx.get("default_pstyle")
    sources += _style_chain(pid, ctx["pstyle"])
    if ctx["docdef"] is not None:
        sources.append(ctx["docdef"])

    def resolve(direct_key, theme_key):
        for src in sources:
            v = src.get(direct_key)
            if v:
                return v
            tv = src.get(theme_key)
            if tv:
                r = _resolve_theme_token(tv, ctx)
                return r if r else _FONT_UNKNOWN
        return _FONT_UNKNOWN
    ea = resolve("eastAsia", "eastAsiaTheme")
    asc = resolve("ascii", "asciiTheme")
    sz = _FONT_UNKNOWN
    for src in sources:
        if src.get("sz") is not None:
            sz = src["sz"]
            break
    return asc, ea, sz


def set_run_fonts(run, p, ascii_font, ea_font, size_pt, changes, idx, ctx):
    eff_asc, eff_ea, eff_sz = effective_run_fonts(run, p, ctx)
    # 字体:仅当解析后的有效中/西文字体与规范不一致时才写入
    if eff_asc != ascii_font or eff_ea != ea_font:
        run.font.name = ascii_font  # 设置 ascii + hAnsi
        run._r.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), ea_font)
        changes.append({"prop": f"run{idx}.字体",
                        "old": f"西文{eff_asc}/中文{eff_ea}",
                        "new": f"西文{ascii_font}/中文{ea_font}"})
    # 字号:仅当有效字号与规范不一致(或无法确定)时才写入
    if eff_sz == _FONT_UNKNOWN or abs(float(eff_sz) - float(size_pt)) > 0.05:
        run.font.size = Pt(size_pt)
        old = "(继承/未知)" if eff_sz == _FONT_UNKNOWN else f"{eff_sz}pt"
        changes.append({"prop": f"run{idx}.字号", "old": old, "new": f"{size_pt}pt"})


def describe_line(pf):
    rule, ls = pf.line_spacing_rule, pf.line_spacing
    if rule == WD_LINE_SPACING.EXACTLY and ls is not None:
        return ("exact", round(ls.pt, 1))
    if rule == WD_LINE_SPACING.SINGLE:
        return ("single", None)
    if rule == WD_LINE_SPACING.MULTIPLE or isinstance(ls, float):
        return ("multiple", ls)
    if ls is None:
        return (None, None)
    return ("other", str(ls))


def apply_para_format(p, s, changes):
    pf = p.paragraph_format
    # 对齐
    want = ALIGN[s["align"]]
    if pf.alignment != want:
        changes.append({"prop": "对齐",
                        "old": ALIGN_NAME.get(pf.alignment, str(pf.alignment or "(继承)")),
                        "new": s["align"]})
        pf.alignment = want
    # 行距(含公式的段落按规范允许按需调整,正文角色时跳过)
    rule, val = s["line"]
    if not (para_has_math(p) and s is spec.BODY):
        cur = describe_line(pf)
        tgt = (rule, float(val) if val else None)
        if cur != tgt:
            changes.append({"prop": "行距", "old": str(cur), "new": str(tgt)})
            if rule == "exact":
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = Pt(val)
            elif rule == "single":
                pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            elif rule == "multiple":
                pf.line_spacing = float(val)
    # 段前段后
    for attr, key, label in (("space_before", "before", "段前"),
                             ("space_after", "after", "段后")):
        cur = getattr(pf, attr)
        cur_pt = round(cur.pt, 1) if cur is not None else None
        if cur_pt != float(s[key]):
            changes.append({"prop": label, "old": f"{cur_pt}磅" if cur_pt is not None
                            else "(继承)", "new": f"{s[key]}磅"})
            setattr(pf, attr, Pt(s[key]))
    # 首行缩进(2汉字符) / 悬挂缩进 —— 用 w:ind 的 firstLineChars/hangingChars,随字号自适应
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if "first_indent_chars" in s:
        n = s["first_indent_chars"] * 100
        cur = ind.get(qn("w:firstLineChars")) if ind is not None else None
        if cur != str(n):
            changes.append({"prop": "首行缩进", "old": f"{cur or '(无)'}(百分字符)",
                            "new": f"{s['first_indent_chars']}字符"})
            if ind is None:
                ind = pPr.makeelement(qn("w:ind"), {})
                pPr.append(ind)
            for a in ("w:firstLine", "w:hanging", "w:hangingChars"):
                if ind.get(qn(a)):
                    del ind.attrib[qn(a)]
            ind.set(qn("w:firstLineChars"), str(n))
    elif "hanging_chars" in s:
        n = int(s["hanging_chars"] * 100)
        cur = ind.get(qn("w:hangingChars")) if ind is not None else None
        if cur != str(n):
            changes.append({"prop": "悬挂缩进", "old": f"{cur or '(无)'}(百分字符)",
                            "new": f"{s['hanging_chars']}字符"})
            if ind is None:
                ind = pPr.makeelement(qn("w:ind"), {})
                pPr.append(ind)
            for a in ("w:firstLine", "w:firstLineChars"):
                if ind.get(qn(a)):
                    del ind.attrib[qn(a)]
            ind.set(qn("w:hangingChars"), str(n))
            ind.set(qn("w:hanging"), str(int(s["hanging_chars"] * 240)))
    elif s["align"] == "center":
        # 居中段落(章标题/题注/公式/表格单元格等):
        #   · 只有"有效缩进非 0"时才把首行/悬挂/左缩进**显式清零**(覆盖样式继承),
        #     否则"先缩进再居中"会让文字/数字整体右偏、不在正中;
        #   · 本就无有效缩进的居中段落**保持原样**——尤其**绝不删除**单元格里已经
        #     显式写好的 firstLine="0",否则反而会暴露段落样式中的首行缩进、造成右偏。
        # 左对齐角色(目录/资料来源等)不走这里,保留其层级左缩进。
        fl, hg, lf = effective_indent(p, _FONT_CTX) if _FONT_CTX is not None else (0, 0, 0)
        if fl or hg or lf:
            if ind is None:
                ind = pPr.makeelement(qn("w:ind"), {})
                pPr.append(ind)
            for a in ("w:hanging", "w:hangingChars", "w:start", "w:startChars"):
                if ind.get(qn(a)):
                    del ind.attrib[qn(a)]
            ind.set(qn("w:firstLineChars"), "0")
            ind.set(qn("w:firstLine"), "0")
            ind.set(qn("w:leftChars"), "0")
            ind.set(qn("w:left"), "0")
            changes.append({"prop": "缩进",
                            "old": f"首行{fl}/悬挂{hg}/左{lf}(含样式继承)",
                            "new": "清零(居中段落不缩进)"})
    else:
        # 非居中段落(左对齐目录/资料来源等):仅清除遗留的直接首行缩进,保留左缩进
        if ind is not None and (ind.get(qn("w:firstLineChars")) or ind.get(qn("w:firstLine"))):
            changes.append({"prop": "首行缩进", "old": "有", "new": "无"})
            for a in ("w:firstLine", "w:firstLineChars"):
                if ind.get(qn(a)):
                    del ind.attrib[qn(a)]


def is_cjk(ch):
    return "\u4e00" <= ch <= "\u9fff" or ch in "。,、;:?!“”‘’()《》—…·"


def apply_run_formats(p, s, changes, ctx):
    for i, run in enumerate(p.runs):
        if run_has_math(run):
            continue
        t = run.text
        if not t.strip():
            continue
        # 纯西文 run 中文字体也按规范设置,不影响显示;混排 run 按中西文各自字体设置
        set_run_fonts(run, p, s["en_font"], s["cn_font"], s["size"], changes, i, ctx)


def cover_section_indices(doc, smap):
    """返回封面/前置区(摘要之前)所占的分节序号集合。

    这些分节有**封面专用页边距**(中文封面 6/6/4/4、书脊 5.5/5.5/1/1、英文封面
    5.5/5/3.6/3.6 cm 等),整改时必须保留,绝不能统一成正文的 3/3/3/3——否则封面
    会整体移位、压缩、错版(模板被搞崩)。判据:某分节若不含任何"内容段"(摘要及
    之后的段落),即视为前置封面区。"""
    cover_end = None
    for e in smap:
        if e.get("role") not in ("cover", "skip"):
            cover_end = e["index"]
            break
    if cover_end is None:
        return set()
    sec_has_content, para_seen, sec_idx = {}, 0, 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            if para_seen >= cover_end:
                sec_has_content[sec_idx] = True
            if child.find(qn("w:pPr") + "/" + qn("w:sectPr")) is not None:
                sec_idx += 1
            para_seen += 1
    return {i for i in range(sec_idx + 1) if not sec_has_content.get(i)}


def apply_section_setup(doc, report, cover_secs=frozenset()):
    for si, sec in enumerate(doc.sections):
        if si in cover_secs:
            continue          # 封面/前置区分节:保留模板专用页边距,完全不动
        changes = []
        for attr, key, label in (("top_margin", "margin_top_cm", "上边距"),
                                 ("bottom_margin", "margin_bottom_cm", "下边距"),
                                 ("left_margin", "margin_left_cm", "左边距"),
                                 ("right_margin", "margin_right_cm", "右边距"),
                                 ("gutter", "gutter_cm", "装订线"),
                                 ("header_distance", "header_cm", "页眉距边界"),
                                 ("footer_distance", "footer_cm", "页脚距边界")):
            cur = getattr(sec, attr)
            cur_cm = round(cur.cm, 2) if cur is not None else None
            want = spec.PAGE[key]
            if cur_cm != want:
                changes.append({"prop": label, "old": f"{cur_cm}cm", "new": f"{want}cm"})
                setattr(sec, attr, Cm(want))
        if changes:
            report.append({"paragraph": f"节{si}", "text": "(页面设置)",
                           "role": "section", "changes": changes})


def _border_el(parent_tag_owner, tag, sz_eighth):
    el = parent_tag_owner.makeelement(qn(f"w:{tag}"), {})
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz_eighth))
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), "000000")
    return el


def apply_three_line_borders(table):
    """三线表:表上下边线1.5磅(sz=12),表头行下线1磅(sz=8),清除其余框线。
    返回修改记录列表。边框 sz 单位为 1/8 磅。"""
    changes = []
    tbl = table._tbl
    tblPr = tbl.tblPr
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    borders.append(_border_el(tblPr, "top", 12))
    borders.append(_border_el(tblPr, "bottom", 12))
    for tag in ("left", "right", "insideH", "insideV"):
        none_el = tblPr.makeelement(qn(f"w:{tag}"), {})
        none_el.set(qn("w:val"), "none")
        none_el.set(qn("w:sz"), "0")
        none_el.set(qn("w:space"), "0")
        borders.append(none_el)
    tblPr.append(borders)
    changes.append({"prop": "表边框", "old": "(原线型)",
                    "new": "三线表:上下1.5磅,内线清除"})
    # 清除单元格级边框,再给表头行下边加 1 磅线
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tb = tcPr.find(qn("w:tcBorders"))
            if tb is not None:
                tcPr.remove(tb)
            if ri == 0:
                tb = tcPr.makeelement(qn("w:tcBorders"), {})
                tb.append(_border_el(tcPr, "bottom", 8))
                tcPr.append(tb)
    changes.append({"prop": "表头下线", "old": "(原线型)", "new": "1磅单线"})
    return changes


def clear_all_table_borders(table):
    """清除表级与单元格级的全部边框(用于图片占位/布局表)。
    这些表只用来约束图片宽度,原本无框/透明;若误当三线表会在 PDF 里出现黑框。
    返回修改记录列表。"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for tag in ("top", "bottom", "left", "right", "insideH", "insideV"):
        e = tblPr.makeelement(qn(f"w:{tag}"), {})
        e.set(qn("w:val"), "nil")
        e.set(qn("w:sz"), "0")
        e.set(qn("w:space"), "0")
        borders.append(e)
    tblPr.append(borders)
    # 单元格级边框也一并清掉
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tb = tcPr.find(qn("w:tcBorders"))
            if tb is not None:
                tcPr.remove(tb)
    return [{"prop": "表边框", "old": "(原线型)", "new": "全部清除(图片布局表)"}]


def extract_all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    ap.add_argument("-m", "--map", required=True, help="structure_map.json(已人工复核)")
    ap.add_argument("-o", "--out", help="输出 docx(整改模式必填)")
    ap.add_argument("-r", "--report", default="format_report.json")
    ap.add_argument("--dry-run", action="store_true", help="只检查不写文件")
    ap.add_argument("--skip-page-setup", action="store_true")
    args = ap.parse_args()

    doc = Document(args.docx)
    global _FONT_CTX
    _FONT_CTX = build_font_context(doc)
    raw = json.loads(Path(args.map).read_text(encoding="utf-8"))
    smap = raw["paragraphs"] if isinstance(raw, dict) else raw
    table_roles = {t["index"]: t.get("role", "three_line")
                   for t in (raw.get("tables", []) if isinstance(raw, dict) else [])}
    if len(smap) != len(doc.paragraphs):
        sys.exit(f"结构图谱段数({len(smap)})与文档段数({len(doc.paragraphs)})不一致,"
                 "请重新运行 classify.py")

    text_before = extract_all_text(doc)
    report = []

    if not args.skip_page_setup:
        apply_section_setup(doc, report, cover_section_indices(doc, smap))

    for entry, p in zip(smap, doc.paragraphs):
        role = entry["role"]
        if role in NO_TOUCH_ROLES or role not in spec.ROLE_SPEC:
            continue
        s = spec.ROLE_SPEC[role]
        changes = []
        apply_para_format(p, s, changes)
        apply_run_formats(p, s, changes, _FONT_CTX)
        if changes:
            report.append({"paragraph": entry["index"],
                           "text": entry["text"], "role": role,
                           "changes": changes})

    # 表格:单元格文字格式 + 三线表边框(role=three_line 时)
    s = spec.TABLE_CELL
    for ti, table in enumerate(doc.tables):
        t_role = table_roles.get(ti, "three_line")
        if t_role == "skip":
            continue
        if t_role == "image_layout":
            # 图片占位/布局表:清掉边框,且不处理单元格文字(里面是图片,无需正文格式)
            bc = clear_all_table_borders(table)
            report.append({"paragraph": f"表{ti+1}", "text": "(图片占位表)",
                           "role": "image_layout_borders", "changes": bc})
            continue
        if t_role == "three_line":
            bc = apply_three_line_borders(table)
            if bc:
                report.append({"paragraph": f"表{ti+1}", "text": "(边框)",
                               "role": "table_borders", "changes": bc})
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    changes = []
                    apply_para_format(p, s, changes)
                    apply_run_formats(p, s, changes, _FONT_CTX)
                    if changes:
                        report.append({"paragraph": f"表{ti+1}单元格({ri},{ci})",
                                       "text": p.text.strip()[:40],
                                       "role": "table_cell", "changes": changes})

    text_after = extract_all_text(doc)
    integrity = "PASS" if text_before == text_after else "FAIL"

    out = {"mode": "dry-run" if args.dry_run else "apply",
           "source": args.docx,
           "content_integrity": integrity,
           "total_paragraphs": len(doc.paragraphs),
           "items_changed": len(report),
           "details": report}
    Path(args.report).write_text(json.dumps(out, ensure_ascii=False, indent=1),
                                 encoding="utf-8")
    print(f"[{out['mode']}] 共 {len(report)} 处需要/已经修改 → {args.report}")
    print(f"内容完整性校验: {integrity}")
    if integrity == "FAIL":
        sys.exit("致命错误:文本内容发生了变化,放弃保存!")

    if not args.dry_run:
        if not args.out:
            sys.exit("整改模式必须指定 -o 输出路径")
        doc.save(args.out)
        print(f"已保存 → {args.out}")


if __name__ == "__main__":
    main()
