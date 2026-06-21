"""make_sample.py — 生成一份格式故意不合规的模拟论文,用于测试整改流水线。"""
import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


def add(doc, text, size=12, cn="楷体", en="Calibri", align=None, bold=False,
        before=None, after=None, line=None, indent=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = en
    r._r.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), cn)
    r.bold = bold
    if align:
        p.paragraph_format.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if line is not None:
        p.paragraph_format.line_spacing = line  # float = 倍数
    if indent:
        p.paragraph_format.first_line_indent = Pt(indent)
    return p


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("-o", "--out", default="sample_thesis.docx",
                    help="输出 docx 路径,默认写到当前目录")
    ap.add_argument("--degree", choices=["博士", "硕士"], default="博士",
                    help="样本学位档次,影响封面学位标识(博硕正文格式相同),默认博士")
    args = ap.parse_args()
    degree = args.degree
    en_degree = "Doctor of Philosophy" if degree == "博士" else "Master of Science"
    en_kind = "Dissertation" if degree == "博士" else "Thesis"

    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)

    C = WD_ALIGN_PARAGRAPH.CENTER
    L = WD_ALIGN_PARAGRAPH.LEFT

    # —— 中文封面(含学位类别行,供 detect_degree.py 识别;封面区由 classify 标为 cover,不整改内容)——
    add(doc, "基于深度学习的工业过程优化方法研究", size=26, cn="中易黑体", align=C, bold=True)
    add(doc, f"(申请清华大学工学{degree}学位论文)", size=16, cn="宋体", align=C)
    add(doc, "培养单位:自动化系", size=16, cn="仿宋", align=C)
    add(doc, "学　　科:控制科学与工程", size=16, cn="仿宋", align=C)
    add(doc, "研究生:张　三", size=16, cn="仿宋", align=C)
    add(doc, "指导教师:李　四 教授", size=16, cn="仿宋", align=C)
    # —— 英文封面(含英文学位行)——
    add(doc, "Research on Deep-Learning-Based Industrial Process Optimization",
        size=20, en="Arial", align=C, bold=True)
    add(doc, f"{en_kind} submitted to Tsinghua University in partial fulfillment of "
             f"the requirement for the degree of {en_degree} in Control Science and Engineering",
        size=12, en="Times New Roman", align=C)
    add(doc, "by Zhang San", size=12, en="Times New Roman", align=C)
    # —— 关于学位论文使用授权的说明(博/硕标题相同,前置签名页)——
    add(doc, "关于学位论文使用授权的说明", size=16, cn="宋体", align=C, bold=True)
    add(doc, "本人完全了解清华大学有关保留、使用学位论文的规定。", size=12, cn="楷体")

    add(doc, "摘要", size=18, cn="宋体", align=C, bold=True, before=6, after=6)
    add(doc, "本文研究了基于深度学习的工业过程优化方法。针对传统方法在高维非线性系统中的局限,提出了一种混合建模框架。", size=12, cn="楷体", line=1.5)
    add(doc, "关键词:深度学习;过程优化;混合建模", size=12, cn="楷体")
    add(doc, "Abstract", size=18, cn="宋体", en="Times New Roman", align=C, bold=True)
    add(doc, "This dissertation studies industrial process optimization based on deep learning.", size=12, en="Calibri")
    add(doc, "Keywords: deep learning; process optimization", size=12, en="Calibri")
    add(doc, "目录", size=18, cn="宋体", align=C, bold=True)
    add(doc, "第1章 引言………………………………………1", size=10.5, cn="宋体")
    add(doc, "1.1 研究背景………………………………………1", size=10.5, cn="宋体")
    add(doc, "第1章 引言", size=15, cn="宋体", align=L, bold=True, before=10, after=10)
    add(doc, "1.1 研究背景", size=14, cn="宋体", bold=True, before=6, after=6)
    add(doc, "随着工业4.0的推进,流程工业对智能优化的需求日益增长。传统的机理建模方法在面对复杂系统时遇到了瓶颈。", size=12, cn="楷体", line=1.5, indent=21)
    add(doc, "1.1.1 国内外研究现状", size=12, cn="宋体", bold=True)
    add(doc, "近年来,数据驱动方法取得了显著进展。", size=12, cn="楷体", line=1.5)
    add(doc, "图1.1 研究框架示意图", size=12, cn="宋体", align=C, before=0, after=0)
    add(doc, "表1-1 各方法性能对比", size=12, cn="宋体", align=C)
    t = doc.add_table(rows=2, cols=3)
    for ri, row in enumerate(t.rows):
        for ci, cell in enumerate(row.cells):
            cell.paragraphs[0].add_run(f"数据{ri}{ci}")
    add(doc, "E = mc² (1-1)", size=12, en="Calibri", align=C)
    add(doc, "参考文献", size=18, cn="宋体", align=C, bold=True)
    add(doc, "[1] 陈登原. 国史旧闻: 第一卷[M]. 北京: 中华书局, 2000: 29.", size=12, cn="楷体")
    add(doc, "[2] Calms R B. Infrared spectroscopic studies[D]. Berkeley: Univ. of California, 1965.", size=12, en="Calibri")
    add(doc, "致谢", size=18, cn="宋体", align=C, bold=True)
    add(doc, "感谢导师的悉心指导。", size=12, cn="楷体", line=1.5)
    add(doc, "声明", size=18, cn="宋体", align=C, bold=True)
    add(doc, "本人郑重声明:所呈交的学位论文是本人在导师指导下独立进行研究工作所取得的成果。", size=12, cn="楷体")

    doc.save(str(out))
    print(f"sample saved → {out}")


if __name__ == "__main__":
    main()
